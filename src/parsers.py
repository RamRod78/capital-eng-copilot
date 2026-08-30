import io
import logging
from typing import Tuple, Optional
import pandas as pd

logger = logging.getLogger(__name__)


def parse_uploaded_file(file_obj, filename: str) -> Tuple[str, Optional[str]]:
    """
    Extract raw text from an uploaded file object based on extension.
    Supports PDF (.pdf), Word (.docx), Excel (.xlsx, .xls), CSV (.csv), and text/markdown (.txt, .md).
    Returns (extracted_text, error_message).
    """
    if not file_obj:
        return "", "File object is empty"

    ext = filename.lower().split(".")[-1] if "." in filename else ""

    try:
        if ext in ("txt", "md", "json"):
            # Plain text / Markdown
            content = file_obj.read()
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="replace"), None
            return str(content), None

        elif ext == "pdf":
            # PDF Extraction using pypdf
            import pypdf
            pdf_reader = pypdf.PdfReader(file_obj)
            extracted_pages = []
            for idx, page in enumerate(pdf_reader.pages, start=1):
                text = page.extract_text() or ""
                if text.strip():
                    extracted_pages.append(f"--- [Page {idx}] ---\n{text.strip()}")
            
            full_text = "\n\n".join(extracted_pages)
            if not full_text.strip():
                return "", "PDF was read but contained no extractable text (it might be a scanned image)."
            return full_text, None

        elif ext in ("docx", "doc"):
            # Word Document Extraction
            import docx
            doc = docx.Document(file_obj)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            
            # Also extract tables if present
            table_texts = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                    if row_text:
                        table_texts.append(row_text)

            full_text = "\n".join(paragraphs)
            if table_texts:
                full_text += "\n\n--- Document Tables ---\n" + "\n".join(table_texts)

            if not full_text.strip():
                return "", "Word document contained no extractable text."
            return full_text, None

        elif ext in ("xlsx", "xls"):
            # Excel Spreadsheet Extraction
            excel_data = pd.read_excel(file_obj, sheet_name=None)
            sheet_texts = []
            for sheet_name, df in excel_data.items():
                sheet_texts.append(f"--- [Sheet: {sheet_name}] ---")
                sheet_texts.append(df.to_string(index=False))
            
            full_text = "\n\n".join(sheet_texts)
            return full_text, None

        elif ext == "csv":
            # CSV Extraction
            df = pd.read_csv(file_obj)
            return df.to_string(index=False), None

        else:
            # Fallback attempt text decode
            content = file_obj.read()
            if isinstance(content, bytes):
                return content.decode("utf-8", errors="replace"), None
            return str(content), None

    except Exception as e:
        logger.error(f"Error parsing file '{filename}': {e}", exc_info=True)
        return "", f"Failed to parse '{filename}': {str(e)}"
